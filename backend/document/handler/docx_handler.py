from docx import Document
from .base import BaseDocumentHandler


class DocxHandler(BaseDocumentHandler):

    async def translate(self, file_path: str, text_translator, target_language: str) -> str:
        doc = Document(file_path)

        # Translate paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    run.text = await text_translator.translate_text(
                        run.text, target_language
                    )

        # Translate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                run.text = await text_translator.translate_text(
                                    run.text, target_language
                                )

        # TODO (optional advanced):
        # - headers
        # - footers
        # - textboxes
        # - footnotes

        output_path = file_path.replace(".docx", f"_{target_language}.docx")
        doc.save(output_path)

        return output_path